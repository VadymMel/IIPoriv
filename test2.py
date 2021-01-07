# # -*- coding: utf8 -*-
#
# from fuzzywuzzy import fuzz
# import docx
#
# doc = docx.Document('LOAN.docx')
# dlinadoc = (len(doc.paragraphs))
# # table = doc.tables[0]
# # table_size = len(table.rows)
# # print(table_size)
# # print(dlinadoc)
# def foo():
#     for table in doc.tables:
#         for row in (table.rows[1], table.rows[-1]):
#             for cell in row.cells:
#                 print(cell.text)
#
#     # for j in range(0, table_size):
#     #     for zx in table.rows[j].cells[0].paragraphs:
# # if not dlinadoc:
# #     pass
# # # elif not table_size:
# # #     pass
# # else:
# #     foo()
#
# foo()
import docx

wordDoc = docx.Document('LOAN.docx')

for table in wordDoc.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)

for parag in wordDoc.paragraphs:
    print(f'paragraf = {parag.text}')

print(len(wordDoc.paragraphs))