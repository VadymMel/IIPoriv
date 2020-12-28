# Import libraries

from gensim.models import doc2vec
from collections import namedtuple
import docx
from pprint import pprint
KEY_COMPARE = ["This is a sentence", "Послідовність виконання та типовий зміст робіт кожного"]
doc = docx.Document('kszi.docx')
# # количество абзацев в документе
dlinadoc = (len(doc.paragraphs))
# Load data

doc1 = ["This is a sentence", "This is another sentence"]

# Transform data (you can add more data preprocessing steps)

docs = []
analyzedDocument = namedtuple('AnalyzedDocument', 'words tags')
# for j in range(0, dlinadoc):
for i, text in enumerate(doc.paragraphs):
    if not text.text:
        pass
    else:
        words = text.text.lower().split()
        tags = [i]
        docs.append(analyzedDocument(words, tags))

# Train model (set min_count = 1, if you want the model to work with the provided example data set)

model = doc2vec.Doc2Vec(docs, vector_size=100, window=300, min_count=1, workers=4)

# Get the vectors

model.docvecs[0]
model.docvecs[1]
pprint(docs)
