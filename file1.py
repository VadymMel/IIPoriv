# -*- coding: utf8 -*-

from fuzzywuzzy import fuzz
import docx
import openpyxl
from pprint import pprint
from openpyxl.styles import Font
from openpyxl import load_workbook
from tkinter.filedialog import askopenfilename, asksaveasfilename

TABLICA = [
    'Параметр',
    'Номер облікового запису',
    'Дата облікового запису',
    'Код та найменування резидента позичальника',
    'Назва основного договору',
    'Номер основного договору',
    'Дата основного договору',
    'Додаткові документи/угоди до основного договору, що є невід’ємними частинами основного договору, а також інші договори/документи, які стосуються реалізації основного договору та здійснення валютних операцій за основним договором [уключаючи документи, які встановлюють/змінюють графік (строки та суми) проведення операцій з одержання/повернення кредиту (позики) та сплати інших платежів позичальником за основним договором]',
    'Код та найменування банку, через який резидент-позичальник здійснює грошові розрахунки (платежі) за основним договором',
    'Загальний обсяг кредиту/позики/поворотної фінансової допомоги/кредитної лінії (сума; валюта/банківський метал)',
    'Граничний строк виконання резидентом-позичальником платежів за договором (кінцева дата згідно з умовами договору)',
    'Додаткова інформація',
    'Відомості про внесення змін до облікового запису (номер, дата)',
    ]


MONTH = {'01': 'січня', '02': 'лютого', '03': 'березня', '04': 'квітня', '05': 'травня', '06': 'червня',
         '07': 'липня', '08': 'серпня', '09': 'вересня', '10': 'жовтня',
         '11': 'листопада', '12': 'грудня'}
VALUTA = ['євро', 'долар', 'гривня']
DATA = []
KOD_NAZ = []
SUM_VAL = []


wb = openpyxl.Workbook()
wb.create_sheet(title='ЗЛ', index=0)
sheet = wb['ЗЛ']
doc = docx.Document('dogovor2.docx')
dlinadoc = (len(doc.paragraphs))
table = doc.tables[0]
table_size = len(table.rows)
KEY_COMPARE = open('KEY_COMPARE.ini', 'r', encoding='utf-8')


for pole_znacenie in range(len(TABLICA)):
    value = TABLICA[pole_znacenie]
    cell = sheet.cell(row=pole_znacenie + 1, column=1)
    cell.value = value

for i in KEY_COMPARE:
    for j in range(0, table_size):
        for zx in table.rows[j].cells[0].paragraphs:
            x = fuzz.partial_ratio(zx.text, i)
            if x > 85:
                if zx.text == ' ':
                    pass
                else:
                    if 'договір позики' in zx.text.lower():
                        value = zx.text[0:(zx.text.find("№"))]
                        cell = sheet.cell(row=5, column=2)
                        cell.value = value
                        value = zx.text[zx.text.find("№") + 1:]
                        cell = sheet.cell(row=6, column=2)
                        cell.value = value
                    for slova in zx.text.replace('"', '').split():
                        '''определение ЕДРПОУ'''
                        try:
                            EDRPOU = int(slova)
                            if len(str(EDRPOU)) == 8:
                                KOD_NAZ.append(f'{slova}; ')
                        except ValueError:
                            pass

                        '''определение даты'''
                        try:
                            datte = int(slova)
                            if 2020 < datte < 2050 or 1 <= datte <= 31:
                                DATA.append(slova)
                        except ValueError:
                            for key, value in MONTH.items():
                                if slova == value:
                                    DATA.append(key)
                                else:
                                    pass
                    if 'ТОВ' in zx.text:
                        '''определение названия'''
                        KOD_NAZ.append(zx.text[zx.text.find('ТОВ'):(zx.text.find(','))])
                    if 'сумі' in zx.text:
                        for slova in zx.text.replace('.', '').replace(',', '.').split():
                            try:
                                summ = float(slova)
                                if summ > 0:
                                    summ = str(summ)
                                    SUM_VAL.append(summ)
                            except ValueError:
                                pass
                            if slova.lower() in VALUTA:
                                SUM_VAL.append(slova.upper())

value = '; '.join(SUM_VAL)
cell = sheet.cell(row=10, column=2)
cell.value = value

value = ''.join(KOD_NAZ)
cell = sheet.cell(row=4, column=2)
cell.value = value

value = '.'.join(DATA)
cell = sheet.cell(row=11, column=2)
cell.value = value

wb.save("23.xlsx")
